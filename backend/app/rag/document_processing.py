"""Safe local validation, extraction, normalization, and deterministic chunking."""

from __future__ import annotations

import io
import re
import unicodedata
import zipfile
from dataclasses import dataclass

from docx import Document
from pypdf import PdfReader


class DocumentProcessingError(ValueError):
    def __init__(self, code: str) -> None:
        self.code = code
        super().__init__(code)


@dataclass(frozen=True)
class ExtractedDocument:
    mime_type: str
    page_count: int | None
    text: str


ALLOWED = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}


def validate_and_extract(
    data: bytes, filename: str, declared_type: str, max_pages: int, max_chars: int
) -> ExtractedDocument:
    extension = filename.lower().rsplit(".", 1)
    extension = f".{extension[-1]}" if len(extension) == 2 else ""
    if extension not in ALLOWED:
        raise DocumentProcessingError("unsupported_document_type")
    expected = ALLOWED[extension]
    if declared_type and declared_type.split(";", 1)[0].lower() not in {
        expected,
        "application/octet-stream",
    }:
        raise DocumentProcessingError("document_mime_mismatch")
    if not data:
        raise DocumentProcessingError("empty_document")
    if extension == ".pdf":
        if not data.startswith(b"%PDF-"):
            raise DocumentProcessingError("document_content_mismatch")
        try:
            reader = PdfReader(io.BytesIO(data), strict=True)
            if reader.is_encrypted:
                raise DocumentProcessingError("encrypted_document")
            if len(reader.pages) > max_pages:
                raise DocumentProcessingError("pdf_page_limit_exceeded")
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except DocumentProcessingError:
            raise
        except Exception:
            raise DocumentProcessingError("malformed_document") from None
        if not text.strip():
            raise DocumentProcessingError("scanned_pdf")
        return ExtractedDocument(
            expected, len(reader.pages), normalize(text, max_chars)
        )
    if extension == ".docx":
        if not data.startswith(b"PK\x03\x04"):
            raise DocumentProcessingError("document_content_mismatch")
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                entries = archive.infolist()
                if (
                    len(entries) > 1000
                    or sum(i.file_size for i in entries) > max_chars * 20
                    or any(i.file_size > 50 * 1024 * 1024 for i in entries)
                ):
                    raise DocumentProcessingError("document_resource_limit")
                if "word/document.xml" not in archive.namelist():
                    raise DocumentProcessingError("malformed_document")
            text = "\n\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)
        except DocumentProcessingError:
            raise
        except Exception:
            raise DocumentProcessingError("malformed_document") from None
        return ExtractedDocument(expected, None, normalize(text, max_chars))
    if data.startswith((b"%PDF-", b"PK\x03\x04")):
        raise DocumentProcessingError("document_content_mismatch")
    try:
        text = data.decode("utf-8-sig")
    except UnicodeDecodeError:
        raise DocumentProcessingError("malformed_document") from None
    return ExtractedDocument(expected, None, normalize(text, max_chars))


def normalize(text: str, max_chars: int) -> str:
    value = (
        unicodedata.normalize("NFC", text)
        .replace("\r\n", "\n")
        .replace("\r", "\n")
        .replace("\x00", "")
    )
    value = "".join(
        char
        for char in value
        if char == "\n" or not unicodedata.category(char).startswith("C")
    )
    value = re.sub(r"[^\S\n]+", " ", value)
    value = re.sub(r" *\n *", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value).strip()
    if not value:
        raise DocumentProcessingError("empty_extracted_text")
    if len(value) > max_chars:
        raise DocumentProcessingError("extracted_text_limit_exceeded")
    return value


def chunk_text(
    text: str, target: int = 1200, maximum: int = 1600, overlap: int = 200
) -> list[str]:
    paragraphs = text.split("\n\n")
    units: list[str] = []
    for paragraph in paragraphs:
        while len(paragraph) > maximum:
            point = paragraph.rfind(" ", 0, maximum + 1)
            point = point if point > 0 else maximum
            units.append(paragraph[:point].strip())
            paragraph = paragraph[point:].strip()
        if paragraph:
            units.append(paragraph)
    chunks: list[str] = []
    current = ""
    for unit in units:
        candidate = unit if not current else f"{current}\n\n{unit}"
        if current and len(candidate) > target:
            chunks.append(current)
            prefix = current[-overlap:] if len(current) > overlap else current
            current = (
                f"{prefix}\n\n{unit}"
                if len(prefix) + len(unit) + 2 <= maximum
                else unit
            )
        else:
            current = candidate
    if current:
        chunks.append(current)
    return chunks
